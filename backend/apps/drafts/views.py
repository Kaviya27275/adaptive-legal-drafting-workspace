from datetime import timedelta

from rest_framework import viewsets, permissions
from rest_framework.parsers import FormParser, MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView
from django.utils import timezone
from django.http import HttpResponse

from apps.compliance.models import DocumentType
from core.permissions import IsLegalRole, IsRoleWriteAllowed
from .models import Draft
from .models import DraftSandbox
from .serializers import DraftSandboxSerializer, DraftSerializer
from .permissions import IsOwnerAdminOrReadOnly
from apps.versions.services import save_version
from .assistant_services import (
    detect_document_type_with_ai,
    drafting_ai_assist,
    extract_text_from_uploaded_file,
    analyze_legal_document,
    generate_draft_from_template,
    list_document_templates,
    render_document_template,
)


class DraftViewSet(viewsets.ModelViewSet):

    serializer_class = DraftSerializer
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed, IsOwnerAdminOrReadOnly]

    def _purge_expired_trash(self):
        cutoff = timezone.now() - timedelta(days=30)
        Draft.objects.filter(deleted_at__lte=cutoff).delete()

    def get_queryset(self):
        self._purge_expired_trash()
        user = self.request.user
        if user.role == "admin":
            return Draft.objects.filter(is_active=True, deleted_at__isnull=True)
        return Draft.objects.filter(created_by=user, is_active=True, deleted_at__isnull=True)

    def perform_create(self, serializer):
        draft = serializer.save(created_by=self.request.user)
        save_version(draft, self.request.user)

    def perform_update(self, serializer):
        draft = serializer.save()
        save_version(draft, self.request.user)

    def perform_destroy(self, instance):
        instance.is_active = False
        instance.deleted_at = timezone.now()
        instance.save(update_fields=["is_active", "deleted_at", "updated_at"])


class DraftUploadAnalyzeView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            return Response({"error": "file is required"}, status=400)
        try:
            content = extract_text_from_uploaded_file(uploaded_file)
        except ValueError as exc:
            return Response({"error": str(exc)}, status=400)

        detected = detect_document_type_with_ai(content)
        return Response(
            {
                "filename": uploaded_file.name,
                "text_excerpt": content[:1500],
                "detected_document_type": detected,
            }
        )


class DraftDocumentAnalysisView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            return Response({"error": "file is required"}, status=400)
        try:
            content = extract_text_from_uploaded_file(uploaded_file)
        except ValueError as exc:
            return Response({"error": str(exc)}, status=400)

        detected = detect_document_type_with_ai(content)
        document_type = detected.get("document_type") if isinstance(detected, dict) else detected
        analysis = analyze_legal_document(content, document_type or "Unknown")

        return Response(
            {
                "filename": uploaded_file.name,
                "content": content,
                "detected_document_type": detected,
                "analysis": analysis,
            }
        )


class DraftTemplateListView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]

    def get(self, request):
        return Response({"templates": list_document_templates()})


class DraftTemplateRenderView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]

    def get(self, request, document_type_id):
        document_type = DocumentType.objects.filter(id=document_type_id).first()
        if not document_type:
            return Response({"error": "Document type not found"}, status=404)
        jurisdiction = request.query_params.get("jurisdiction", "India")
        try:
            content = render_document_template(document_type.name, jurisdiction=jurisdiction)
        except ValueError as exc:
            return Response({"error": str(exc)}, status=400)
        return Response(
            {
                "document_type": document_type.name,
                "jurisdiction": jurisdiction,
                "template_content": content,
            }
        )


class DraftExportView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def _get_draft(self, request, draft_id):
        if request.user.role == "admin":
            return Draft.objects.filter(id=draft_id, is_active=True, deleted_at__isnull=True).first()
        return Draft.objects.filter(
            id=draft_id, created_by=request.user, is_active=True, deleted_at__isnull=True
        ).first()

    def get(self, request, draft_id):
        draft = self._get_draft(request, draft_id)
        if not draft:
            return Response({"error": "Draft not found"}, status=404)

        fmt = (request.query_params.get("file_format") or request.query_params.get("format") or "txt").lower()
        title = (draft.title or "draft").strip()
        content = draft.content or ""

        if fmt == "pdf":
            try:
                from reportlab.pdfgen import canvas
            except ImportError:
                return Response({"error": "PDF export support is not installed."}, status=400)
            response = HttpResponse(content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="{title}.pdf"'
            pdf = canvas.Canvas(response)
            y = 800
            pdf.setFont("Helvetica", 11)
            for line in content.splitlines() or [""]:
                if y <= 40:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 11)
                    y = 800
                pdf.drawString(40, y, line[:120])
                y -= 14
            pdf.save()
            return response

        if fmt == "doc":
            try:
                import pypandoc
            except ImportError:
                return Response({"error": "DOC export support is not installed."}, status=400)
            try:
                rtf = pypandoc.convert_text(content or "", "rtf", format="markdown")
            except OSError:
                return Response({"error": "Pandoc is required for DOC export."}, status=400)
            response = HttpResponse(rtf, content_type="application/rtf")
            response["Content-Disposition"] = f'attachment; filename="{title}.doc"'
            return response

        if fmt == "docx":
            try:
                from docx import Document as DocxDocument
            except ImportError:
                return Response({"error": "DOCX export support is not installed."}, status=400)
            document = DocxDocument()
            document.add_heading(title, level=1)
            for line in content.splitlines():
                document.add_paragraph(line)
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response["Content-Disposition"] = f'attachment; filename="{title}.docx"'
            document.save(response)
            return response

        response = HttpResponse(content, content_type="text/plain")
        response["Content-Disposition"] = f'attachment; filename="{title}.txt"'
        return response


class DraftTrashListView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def get(self, request):
        cutoff = timezone.now() - timedelta(days=30)
        Draft.objects.filter(deleted_at__lte=cutoff).delete()
        if request.user.role == "admin":
            items = (
                Draft.objects.filter(deleted_at__isnull=False)
                .select_related("document_type")
                .order_by("-deleted_at")
            )
        else:
            items = (
                Draft.objects.filter(created_by=request.user, deleted_at__isnull=False)
                .select_related("document_type")
                .order_by("-deleted_at")
            )
        return Response(
            [
                {
                    "id": item.id,
                    "title": item.title,
                    "document_type": item.document_type_id,
                    "document_type_name": item.document_type.name if item.document_type_id else "",
                    "deleted_at": item.deleted_at,
                    "updated_at": item.updated_at,
                }
                for item in items
            ]
        )


class DraftTrashRestoreView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def post(self, request, draft_id):
        if request.user.role == "admin":
            draft = Draft.objects.filter(id=draft_id, deleted_at__isnull=False).first()
        else:
            draft = Draft.objects.filter(
                id=draft_id, created_by=request.user, deleted_at__isnull=False
            ).first()
        if not draft:
            return Response({"error": "Draft not found in trash."}, status=404)

        conflict = Draft.objects.filter(
            created_by=draft.created_by,
            title__iexact=draft.title,
            deleted_at__isnull=True,
            is_active=True,
        ).exists()
        if conflict:
            return Response(
                {"error": "A draft with the same file name already exists."},
                status=400,
            )

        draft.deleted_at = None
        draft.is_active = True
        draft.save(update_fields=["deleted_at", "is_active", "updated_at"])
        return Response({"message": "Draft restored."})


class DraftTrashDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def delete(self, request, draft_id):
        if request.user.role == "admin":
            draft = Draft.objects.filter(id=draft_id, deleted_at__isnull=False).first()
        else:
            draft = Draft.objects.filter(
                id=draft_id, created_by=request.user, deleted_at__isnull=False
            ).first()
        if not draft:
            return Response({"error": "Draft not found in trash."}, status=404)
        draft.delete()
        return Response({"message": "Draft permanently deleted."})


class DraftTrashDeleteAllView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def delete(self, request):
        if request.user.role == "admin":
            Draft.objects.filter(deleted_at__isnull=False).delete()
        else:
            Draft.objects.filter(created_by=request.user, deleted_at__isnull=False).delete()
        return Response({"message": "Trash cleared."})


class DraftAIAssistView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]

    def _get_draft(self, request, draft_id):
        if request.user.role == "admin":
            return Draft.objects.filter(id=draft_id, is_active=True, deleted_at__isnull=True).first()
        return Draft.objects.filter(
            id=draft_id, created_by=request.user, is_active=True, deleted_at__isnull=True
        ).first()

    def post(self, request):
        draft_id = request.data.get("draft_id")
        instruction = request.data.get("instruction", "")
        override_text = request.data.get("text")
        if not draft_id:
            return Response({"error": "draft_id is required"}, status=400)
        draft = self._get_draft(request, draft_id)
        if not draft:
            return Response({"error": "Draft not found"}, status=404)
        assist = drafting_ai_assist(override_text or draft.content, instruction)
        return Response(
            {
                "draft_id": str(draft.id),
                "assistant_output": assist,
                "note": "Draft content is not automatically modified.",
            }
        )


class DraftAIDraftView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]

    def post(self, request):
        document_type_id = request.data.get("document_type_id")
        key_terms = request.data.get("key_terms", [])
        jurisdiction = request.data.get("jurisdiction", "India")
        if not document_type_id:
            return Response({"error": "document_type_id is required"}, status=400)
        document_type = DocumentType.objects.filter(id=document_type_id).first()
        if not document_type:
            return Response({"error": "Document type not found"}, status=404)
        if isinstance(key_terms, str):
            key_terms = [term.strip() for term in key_terms.split(",") if term.strip()]

        content = generate_draft_from_template(
            document_type.name,
            key_terms=key_terms,
            jurisdiction=jurisdiction,
        )

        return Response(
            {
                "document_type": document_type.name,
                "jurisdiction": jurisdiction,
                "key_terms": key_terms,
                "content": content,
            }
        )


class DraftSandboxStartView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def _get_draft(self, request, draft_id):
        if request.user.role == "admin":
            return Draft.objects.filter(id=draft_id, is_active=True, deleted_at__isnull=True).first()
        return Draft.objects.filter(
            id=draft_id, created_by=request.user, is_active=True, deleted_at__isnull=True
        ).first()

    def post(self, request, draft_id):
        draft = self._get_draft(request, draft_id)
        if not draft:
            return Response({"error": "Draft not found"}, status=404)
        sandbox = DraftSandbox.objects.create(
            draft=draft,
            created_by=request.user,
            sandbox_content=draft.content or "",
        )
        return Response(
            {
                "message": "Sandbox started",
                "sandbox": DraftSandboxSerializer(sandbox).data,
            }
        )


class DraftSandboxUpdateView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def _get_sandbox(self, request, sandbox_id):
        if request.user.role == "admin":
            return DraftSandbox.objects.filter(id=sandbox_id, is_active=True).first()
        return DraftSandbox.objects.filter(
            id=sandbox_id,
            created_by=request.user,
            is_active=True,
        ).first()

    def patch(self, request, sandbox_id):
        sandbox = self._get_sandbox(request, sandbox_id)
        if not sandbox:
            return Response({"error": "Sandbox not found"}, status=404)
        sandbox.sandbox_content = request.data.get("sandbox_content", sandbox.sandbox_content)
        sandbox.save(update_fields=["sandbox_content", "updated_at"])
        return Response({"sandbox": DraftSandboxSerializer(sandbox).data})


class DraftSandboxAssistView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsLegalRole]

    def _get_sandbox(self, request, sandbox_id):
        if request.user.role == "admin":
            return DraftSandbox.objects.filter(id=sandbox_id, is_active=True).first()
        return DraftSandbox.objects.filter(
            id=sandbox_id,
            created_by=request.user,
            is_active=True,
        ).first()

    def post(self, request, sandbox_id):
        sandbox = self._get_sandbox(request, sandbox_id)
        if not sandbox:
            return Response({"error": "Sandbox not found"}, status=404)
        instruction = request.data.get("instruction", "")
        assist = drafting_ai_assist(sandbox.sandbox_content, instruction)
        return Response(
            {
                "sandbox_id": sandbox.id,
                "assistant_output": assist,
                "note": "Sandbox content is not automatically modified.",
            }
        )


class DraftSandboxCommitView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def _get_sandbox(self, request, sandbox_id):
        if request.user.role == "admin":
            return DraftSandbox.objects.select_related("draft").filter(id=sandbox_id, is_active=True).first()
        return (
            DraftSandbox.objects.select_related("draft")
            .filter(id=sandbox_id, created_by=request.user, is_active=True)
            .first()
        )

    def post(self, request, sandbox_id):
        sandbox = self._get_sandbox(request, sandbox_id)
        if not sandbox:
            return Response({"error": "Sandbox not found"}, status=404)
        draft = sandbox.draft
        draft.content = sandbox.sandbox_content
        draft.save(update_fields=["content", "updated_at"])
        save_version(draft, request.user)
        sandbox.is_active = False
        sandbox.save(update_fields=["is_active", "updated_at"])
        return Response({"message": "Sandbox committed to draft", "draft_id": draft.id})


class DraftSandboxDiscardView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsRoleWriteAllowed]

    def _get_sandbox(self, request, sandbox_id):
        if request.user.role == "admin":
            return DraftSandbox.objects.filter(id=sandbox_id, is_active=True).first()
        return DraftSandbox.objects.filter(
            id=sandbox_id,
            created_by=request.user,
            is_active=True,
        ).first()

    def post(self, request, sandbox_id):
        sandbox = self._get_sandbox(request, sandbox_id)
        if not sandbox:
            return Response({"error": "Sandbox not found"}, status=404)
        sandbox.is_active = False
        sandbox.save(update_fields=["is_active", "updated_at"])
        return Response({"message": "Sandbox discarded"})
